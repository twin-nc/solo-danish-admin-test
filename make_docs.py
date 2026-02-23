"""
Generates the plain-language project documentation as a Word document.
Run with:  venv/Scripts/python make_docs.py
Output:    docs/Project Overview.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

os.makedirs("docs", exist_ok=True)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)


# ── Helpers ───────────────────────────────────────────────────────────────────

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # dark blue
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11) if p.runs else None
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2E74B5")
        tcPr.append(shd)
    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = val
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()  # spacing after table


def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # light grey background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


# ═══════════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════════

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Danish Tax Administration Platform")
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run("Registration Backend — Plain-Language Project Overview")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run("February 2026").font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. WHAT ARE WE BUILDING?
# ═══════════════════════════════════════════════════════════════════════════════

h1("1. What Are We Building?")

body(
    "We are building the backend of a tax registration system for Danish businesses. "
    "A backend is the part of a software system that runs on a server and does the "
    "actual work — storing data, applying rules, and responding to requests. "
    "You never see the backend directly; it is the engine under the hood."
)
doc.add_paragraph()
body(
    "The specific job of this backend: allow business entities to register themselves "
    "with the tax authority. A user (or a web form) sends their company information, "
    "and the backend saves it, validates it, and hands back a confirmation."
)
doc.add_paragraph()
body(
    "This is the first of several planned modules. Tax Filing and Tax Assessment will "
    "follow later, built on top of the same foundation."
)


# ═══════════════════════════════════════════════════════════════════════════════
# 2. THE FIVE THINGS IT CAN DO
# ═══════════════════════════════════════════════════════════════════════════════

h1("2. The Five Things It Can Do")

body(
    "The backend exposes five \"endpoints\" — think of them as five doors you can knock on. "
    "Each door accepts a specific type of request and gives back a specific response."
)
doc.add_paragraph()

add_table(
    headers=["#", "Request", "What It Does"],
    rows=[
        ("1", "GET  /health",                       "Checks whether the server is running. Returns: status: ok."),
        ("2", "POST /api/v1/parties",               "Registers a new business entity. Returns the saved record with a unique ID."),
        ("3", "GET  /api/v1/parties/{id}",          "Looks up a previously registered business by its ID."),
        ("4", "POST /api/v1/parties/{id}/roles",    "Assigns a business role to a registered entity (e.g. Importer, Farm)."),
        ("5", "GET  /api/v1/parties/{id}/roles",    "Lists all roles assigned to a registered entity."),
    ],
    col_widths=[0.3, 2.3, 3.5],
)

body(
    "POST means \"send data to create something new\". "
    "GET means \"retrieve something that already exists\"."
)


# ═══════════════════════════════════════════════════════════════════════════════
# 3. KEY VOCABULARY
# ═══════════════════════════════════════════════════════════════════════════════

h1("3. Key Vocabulary")

body(
    "The system uses specific terms that come from the official data model. "
    "Here is what each one means in plain language:"
)
doc.add_paragraph()

add_table(
    headers=["Term", "Plain meaning"],
    rows=[
        ("Party",                   "A registered business entity — a company, farm, restaurant, etc."),
        ("Party Role",              "The business type/role that entity plays in the tax system (e.g. Importer, Livestock Farmer)."),
        ("Identifier (TIN)",        "Tax Identification Number — the unique ID the tax authority assigns to the business."),
        ("Classification",          "Two labels: (1) the size of the business (Small/Medium/Large) and (2) its economic sector."),
        ("State",                   "The current status of the entity — for now, always IN_BUSINESS."),
        ("Contact",                 "The business email address on file."),
        ("Name",                    "The legal name of the business (and any trading aliases)."),
        ("Eligible Identifier",     "Which tax ID is linked to a specific role assignment."),
        ("Eligible Contact",        "Which email is linked to a specific role assignment."),
    ],
    col_widths=[2.0, 4.1],
)


# ═══════════════════════════════════════════════════════════════════════════════
# 4. VALID VALUES
# ═══════════════════════════════════════════════════════════════════════════════

h1("4. Valid Field Values")

body(
    "Certain fields only accept specific values. These come from the official classification "
    "lists in the documentation."
)
doc.add_paragraph()

add_table(
    headers=["Field", "Allowed Values"],
    rows=[
        ("Identifier type",           "TIN  (Tax Identification Number — the only type used)"),
        ("Party type",                "ORGADM1  (Organisation — the only type used)"),
        ("Party role type",           "BUSINSSDM1  (Business — the only type used)"),
        ("Classification type",       "BUSINESS_SIZE  or  ECONOMIC_ACTIVITIES"),
        ("Business size",             "SMALL,  MEDIUM,  LARGE"),
        ("Party state",               "IN_BUSINESS"),
        ("Economic sector (sector)",  "PIZZERIA, CAFE, FINE_DINE_MEAT, FINE_DINE_VEGETARIAN,\n"
                                      "CLEANING_SERVICES, DISTRIBUTORS, MEAT_PROCESSORS,\n"
                                      "FOOD_PROCESSORS, LIVESTOCK_FARMERS, CROP_FARMERS,\n"
                                      "IMPORTER, CUSTOMER"),
    ],
    col_widths=[2.2, 3.9],
)


# ═══════════════════════════════════════════════════════════════════════════════
# 5. HOW THE DATA IS STORED
# ═══════════════════════════════════════════════════════════════════════════════

h1("5. How the Data Is Stored")

body(
    "All data lives in a PostgreSQL database — think of it as a very reliable, structured "
    "set of spreadsheets that never loses data. Each \"spreadsheet\" is called a table."
)
doc.add_paragraph()
body("When a business registers, information is spread across 10 related tables:")
doc.add_paragraph()

add_table(
    headers=["Table", "What it stores"],
    rows=[
        ("parties",                         "One row per registered business (the parent record)."),
        ("party_identifiers",               "Their TIN (tax number). One business can have multiple."),
        ("party_classifications",           "Size label and sector label. Typically two rows per business."),
        ("party_states",                    "Their current status (IN_BUSINESS)."),
        ("party_contacts",                  "Email addresses on file."),
        ("party_names",                     "Legal name and any aliases."),
        ("party_roles",                     "The business roles assigned to them."),
        ("party_role_states",               "Status of each role."),
        ("party_role_eligible_identifiers", "Which tax number is linked to which role."),
        ("party_role_eligible_contacts",    "Which email is linked to which role."),
    ],
    col_widths=[2.8, 3.3],
)

body(
    "Every row in every table has a UUID — a random unique ID (like a128-character fingerprint) "
    "that guarantees no two records are ever confused with each other. "
    "Every row also automatically records when it was created and last updated."
)


# ═══════════════════════════════════════════════════════════════════════════════
# 6. HOW THE CODE IS ORGANISED
# ═══════════════════════════════════════════════════════════════════════════════

h1("6. How the Code Is Organised — The Four Layers")

body(
    "The code is split into four layers. Each layer has exactly one job and only "
    "communicates with the layer directly below it. This keeps the code clean and "
    "makes it easy to change one part without breaking the others."
)
doc.add_paragraph()

add_table(
    headers=["Layer", "Job", "Rule"],
    rows=[
        ("Router",      "Receives the HTTP request from the outside world. Passes it to the Service.",         "Never touches the database directly."),
        ("Service",     "Applies business rules. Decides what is allowed and what should happen.",              "Never builds an HTTP response. Never writes SQL."),
        ("Repository",  "Reads from and writes to the database.",                                              "Contains no business logic — just data access."),
        ("Models",      "Defines the exact shape of every database table.",                                    "Pure definitions — no logic at all."),
    ],
    col_widths=[1.3, 2.8, 2.0],
)

body("A request flows in one direction only:")
doc.add_paragraph()
add_code_block(
    "HTTP Request\n"
    "   → Router  (parse the request)\n"
    "       → Service  (apply rules)\n"
    "           → Repository  (read/write database)\n"
    "               → Database  (PostgreSQL)"
)


# ═══════════════════════════════════════════════════════════════════════════════
# 7. THE EVENT SYSTEM
# ═══════════════════════════════════════════════════════════════════════════════

h1("7. The Event System — Announcing What Happened")

body(
    "When something significant happens, the system publishes an announcement called a "
    "domain event. Think of it like ringing a bell. Any part of the system that cares "
    "about that event can listen for the bell and react — without the part that rang it "
    "needing to know who is listening."
)
doc.add_paragraph()

h2("Current events")

bullet("PartyRegistered", bold_prefix="PartyRegistered  —  ")
p = doc.paragraphs[-1]
p.clear()
p.style = doc.styles["List Bullet"]
run = p.add_run("PartyRegistered")
run.bold = True
p.add_run("  —  fired when a business entity is successfully registered.")

p2 = doc.add_paragraph(style="List Bullet")
run2 = p2.add_run("PartyRoleAssigned")
run2.bold = True
p2.add_run("  —  fired when a role is assigned to a registered entity.")

doc.add_paragraph()

h2("What happens right now")
body(
    "Each event causes a log message to be written — a timestamped note in the server's "
    "activity log, e.g.: \"Party registered — id=abc-123  type=ORGADM1  tin=1676-123456789\"."
)
doc.add_paragraph()

h2("Why this matters for the future")
body(
    "When the Tax Filing module is built, it can simply subscribe to PartyRegistered. "
    "The moment a business registers, Tax Filing is automatically notified and can "
    "prepare their filing account — without the Registration module needing any changes."
)
doc.add_paragraph()

h2("The swappable bus")
body(
    "The event system is designed so that the \"carrier\" can be swapped out. "
    "Right now events travel in-process (no extra infrastructure). "
    "When the system grows, a message broker like RabbitMQ or Kafka can be plugged in "
    "by changing a single line in the main configuration file. "
    "Nothing else in the codebase changes."
)


# ═══════════════════════════════════════════════════════════════════════════════
# 8. EXAMPLE: WHAT HAPPENS WHEN A BUSINESS REGISTERS
# ═══════════════════════════════════════════════════════════════════════════════

h1("8. Step-by-Step: What Happens When a Business Registers")

body("Imagine a pizza restaurant wants to register. Here is the exact sequence:")
doc.add_paragraph()

steps = [
    ("Step 1 — Send the request",
     "A web form (or another system) sends a POST request to /api/v1/parties containing the "
     "company's name, TIN, size, sector, email, and current state."),
    ("Step 2 — Router receives it",
     "The Router layer unpacks the request and hands the data to the Service layer."),
    ("Step 3 — Service validates it",
     "The Service layer checks that the data is valid (correct field types, required fields present)."),
    ("Step 4 — Repository saves it",
     "The Repository layer writes one row to the parties table and one row each to "
     "party_identifiers, party_classifications (×2), party_states, party_contacts, and party_names."),
    ("Step 5 — Event is published",
     "The Service publishes a PartyRegistered event. The event handler logs: "
     "\"Party registered — Pizza Palace ApS, TIN 1676-999888777\"."),
    ("Step 6 — Response sent back",
     "The Router sends a 201 Created response back to the caller, containing the full saved "
     "record including the newly assigned UUID."),
]

for title_text, detail_text in steps:
    p = doc.add_paragraph(style="List Number")
    run_t = p.add_run(title_text + "  ")
    run_t.bold = True
    p.add_run(detail_text)

doc.add_paragraph()
body("Assigning a role works the same way, with one extra rule:")
doc.add_paragraph()

p = doc.add_paragraph(style="List Bullet")
run_rule = p.add_run("Ownership check:  ")
run_rule.bold = True
p.add_run(
    "Before saving the role, the Service verifies that every identifier and contact "
    "listed in the role actually belongs to this business. If a foreign ID is supplied "
    "(from a different company), the request is rejected with a 400 Bad Request error."
)


# ═══════════════════════════════════════════════════════════════════════════════
# 9. THE FILES
# ═══════════════════════════════════════════════════════════════════════════════

h1("9. The Files and What They Do")

body("The project is organised into folders. Here is what each one contains:")
doc.add_paragraph()

add_table(
    headers=["File / Folder", "Plain-language purpose"],
    rows=[
        (".env",                        "Secret configuration file. Contains the database password. Never shared publicly."),
        (".env.example",                "A safe template showing what .env should look like, without real passwords."),
        ("requirements.txt",            "The shopping list of all external tools the code depends on."),
        ("ARCHITECTURE.md",             "The technical design document (aimed at developers)."),
        ("app/main.py",                 "The front door. Starts the server and wires all the pieces together."),
        ("app/config.py",               "Reads the .env file and makes the settings available to the rest of the code."),
        ("app/db/session.py",           "Manages the connection to the PostgreSQL database."),
        ("app/models/",                 "Defines the shape of every database table (10 tables total)."),
        ("app/schemas/",                "Defines the shape of every API message — what data comes in and goes out."),
        ("app/repositories/",           "All code that reads from or writes to the database."),
        ("app/services/",               "All business rules and logic."),
        ("app/events/",                 "The event/announcement system: definitions, the bus, and the handlers."),
        ("app/routers/",                "The 5 API endpoints that the outside world can call."),
        ("alembic/",                    "Instructions for building the database tables. Run once before first use."),
        ("tests/",                      "Automated tests. 14 tests covering all 5 API endpoints."),
        ("TESTING.md",                  "Developer guide explaining the test suite, how isolation works, and how to run tests."),
    ],
    col_widths=[2.3, 3.8],
)


# ═══════════════════════════════════════════════════════════════════════════════
# 10. TECH STACK
# ═══════════════════════════════════════════════════════════════════════════════

h1("10. Technology Choices — and Why")

add_table(
    headers=["Technology", "What it is", "Why we use it"],
    rows=[
        ("Python",              "Programming language",         "Widely used; strong libraries for data and tax systems."),
        ("FastAPI",             "Web framework",                "Fast, modern, auto-generates interactive API documentation."),
        ("PostgreSQL",          "Database",                     "Reliable, relational, excellent support for UUIDs."),
        ("SQLAlchemy 2.0",      "Database toolkit (ORM)",       "The standard Python tool for working with relational databases."),
        ("Alembic",             "Database migration tool",      "Safely builds and updates database tables as the code evolves."),
        ("Pydantic v2",         "Data validation library",      "Ensures all incoming and outgoing data has exactly the right shape."),
        ("InMemoryEventBus",    "Event system (current)",       "Runs entirely in-process — no extra infrastructure needed to start."),
        ("RabbitMQ / Kafka",    "Event system (future option)", "Can be dropped in later for large-scale, distributed deployments."),
    ],
    col_widths=[1.8, 1.8, 2.5],
)


# ═══════════════════════════════════════════════════════════════════════════════
# 11. HOW TO RUN IT
# ═══════════════════════════════════════════════════════════════════════════════

h1("11. How to Run It")

body("Once PostgreSQL is installed and running, three commands start the system:")
doc.add_paragraph()

h2("Step 1 — Create the database")
body("Open a terminal and run:")
add_code_block('psql -U postgres -c "CREATE DATABASE vatri_db;"')

h2("Step 2 — Build the database tables")
body("This runs the migration — it builds all 10 tables from the instructions in the alembic/ folder:")
add_code_block("venv\\Scripts\\alembic upgrade head")

h2("Step 3 — Start the server")
add_code_block("venv\\Scripts\\uvicorn app.main:app --reload")

body(
    "The server is now running. Open a browser and go to "
    "http://localhost:8000/docs to see the interactive API explorer "
    "(Swagger UI) — you can test all 5 endpoints directly from the browser."
)
doc.add_paragraph()

h2("Quick verification checklist")
checks = [
    "GET /health  →  {\"status\": \"ok\"}",
    "POST /api/v1/parties with valid data  →  201 response with a UUID",
    "GET /api/v1/parties/{id}  →  200 response with the full record",
    "POST /api/v1/parties/{id}/roles  →  201 response with the role",
    "POST /api/v1/parties/{id}/roles with a foreign identifier  →  400 error",
    "Server log shows event messages when actions are performed",
    "Swagger UI at /docs lists all 5 endpoints",
]
for c in checks:
    doc.add_paragraph(c, style="List Bullet")


# ═══════════════════════════════════════════════════════════════════════════════
# 12. AUTOMATED TESTS
# ═══════════════════════════════════════════════════════════════════════════════

h1("12. Automated Tests")

body(
    "The project includes 14 automated tests that verify every endpoint works "
    "correctly. These tests run against a dedicated test database and are designed "
    "to be run by any developer at any time to confirm the system is healthy."
)
doc.add_paragraph()

h2("What the tests check")

add_table(
    headers=["Area", "Tests", "What is verified"],
    rows=[
        ("Party Registration",  "7 tests",  "Registering a business returns 201, the response contains all fields, "
                                            "the data can be retrieved afterwards, and looking up an unknown ID returns 404."),
        ("Role Assignment",     "7 tests",  "Assigning a role returns 201, the response shape is correct, roles appear "
                                            "in the list endpoint, and invalid identifiers or contacts are rejected with 400."),
    ],
    col_widths=[1.8, 0.8, 3.5],
)

h2("How test isolation works")
body(
    "Every test is wrapped in a database transaction that is automatically rolled "
    "back when the test finishes. This means each test starts with a completely clean "
    "database — no leftover data from previous tests — without needing to rebuild or "
    "truncate tables between runs. The test database (vatri_test_db) is separate from "
    "the development database (vatri_db) and is created automatically."
)
doc.add_paragraph()

h2("Running the tests")
body("From a terminal in the project folder:")
add_code_block("venv\\Scripts\\activate\npytest tests/ -v")
body("All 14 tests should pass in under one second.")


# ═══════════════════════════════════════════════════════════════════════════════
# 13. WHAT COMES NEXT
# ═══════════════════════════════════════════════════════════════════════════════

h1("13. What Comes Next")

body(
    "This backend is the first module. It is designed so that additional modules slot "
    "in cleanly without changing what is already built."
)
doc.add_paragraph()

add_table(
    headers=["Module", "Status", "What it will do"],
    rows=[
        ("Entity Registration",  "Complete",  "Businesses register themselves. Covered by this document."),
        ("Automated Tests",      "Complete",  "14 tests covering all 5 API endpoints. Run with: pytest tests/ -v"),
        ("Tax Filing",           "Planned",   "Registered businesses submit their VAT and tax returns."),
        ("Tax Assessment",       "Planned",   "The authority evaluates the filings and issues assessments."),
        ("Frontend / UI",        "Planned",   "A web interface so staff and businesses interact via a browser, not raw API calls."),
    ],
    col_widths=[1.8, 1.1, 3.2],
)

body(
    "The Tax Filing module will automatically receive a notification (via the event system) "
    "the moment a new business registers — without the Registration module needing to change."
)


# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════

output_path = "docs/Project Overview.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
